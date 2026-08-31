from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\Users\wface\Desktop\Lutealark\Lutealark_呼吸训练与情绪记录功能说明.docx"

INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6470"
CALLOUT = "F4F6F9"


def set_font(run, name="Calibri", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), size=11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), size=11)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.left_indent = Inches(0.22)
    shade_paragraph(p, "F2F4F7")
    set_font(p.add_run(text), name="Consolas", size=9, color="2F4F4F")
    return p


def add_callout(doc, title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    shade_paragraph(p, CALLOUT)
    run = p.add_run(title + "\n")
    set_font(run, size=11, color=DARK_BLUE, bold=True)
    set_font(p.add_run(text), size=11, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def configure(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(4)
    set_font(footer.add_run("Lutealark | 呼吸训练与情绪记录功能说明"), size=9, color=MUTED)


def build_document():
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("Lutealark：呼吸训练与情绪记录功能说明")
    set_font(title_run, size=24, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_font(subtitle.add_run("功能机制、数据处理规则与端到端调试指南"), size=11, color=MUTED)

    add_callout(
        doc,
        "文档范围",
        "本说明覆盖当前 Demo 已实现的呼吸训练、每日情绪记录、最近 7 天历史摘要，以及前端—后端—OpenTrek 平台的调试路径。每日记录目前只保存在浏览器本地，不依赖数据库。",
    )

    add_heading(doc, "1. 呼吸训练", 1)
    add_heading(doc, "1.1 功能目标", 2)
    doc.add_paragraph("当用户表达焦虑、紧绷或主动提出想做呼吸训练时，系统提供温和邀请，并在用户确认后打开前端训练界面。平台负责理解与路由，前端负责实际训练体验。")

    add_heading(doc, "1.2 多轮确认机制", 2)
    add_code(doc, "用户表达焦虑或呼吸需求\n→ 平台回复并发出邀请\n→ cache_set('pending_action', 'breathing')\n→ 用户下一句回答\n→ 动作确认识别\n→ 清除缓存并返回最终动作")
    add_bullet(doc, "pending_action 是平台脚本缓存中的短暂状态，用来准确理解“好”在同意什么。")
    add_bullet(doc, "它不是长期记忆，不负责保存用户的健康、周期或情绪历史。")
    add_bullet(doc, "当用户同意、拒绝或改聊新话题时，工作流都会清除该状态，避免后续消息误判。")

    add_heading(doc, "1.3 用户回答的路由结果", 2)
    add_bullet(doc, "“好 / 可以 / 要” → 清除 pending_action，结果渲染返回 action = open_breathing。")
    add_bullet(doc, "“不要 / 不用了” → 清除 pending_action，进入拒绝后的普通陪伴回复。")
    add_bullet(doc, "“这是什么？” → 解释呼吸训练，保留 pending_action，用户仍可继续确认。")
    add_bullet(doc, "“我想做呼吸训练” → 作为新需求重新识别，即使用户之前拒绝过也可以直接进入呼吸路径。")

    add_heading(doc, "1.4 前端动作机制", 2)
    doc.add_paragraph("平台返回的 metadata.action 决定前端是否显示入口。前端只识别最终确认后的 open_breathing，不会因为普通邀请就直接跳转。")
    add_code(doc, "offer_breathing  → 仅代表正在邀请，不直接显示训练入口\nopen_breathing   → 用户已经确认，前端显示“开始呼吸训练”按钮")
    add_bullet(doc, "点击“开始呼吸训练”后进入呼吸训练页。")
    add_bullet(doc, "训练页提供呼吸节奏、倒计时、进度、暂停、继续与停止。")

    add_heading(doc, "1.5 呼吸训练调试", 2)
    add_number(doc, "在前端聊天输入“我有点焦虑”。")
    add_number(doc, "确认 Agent 提出温和呼吸邀请后，回复“好”。")
    add_number(doc, "确认聊天气泡下方出现“开始呼吸训练”按钮，并点击进入训练页。")
    add_number(doc, "检查计时、呼吸阶段、暂停、继续和停止是否正常。")
    add_number(doc, "在平台 Trace 中确认：情绪支持/呼吸邀请 → 记录待确认动作 → 打开呼吸训练回复 → 结果渲染。")
    add_code(doc, '{ "intent": "emotion_support", "action": "open_breathing" }')

    add_heading(doc, "2. 每日情绪与状态记录", 1)
    add_heading(doc, "2.1 记录内容与本地保存", 2)
    doc.add_paragraph("用户在“周期状态”页主动填写当日状态。当前版本不会由 Agent 自动创建或替用户保存记录。")
    add_code(doc, '{\n  "date": "2026-07-31",\n  "energy": 2,\n  "mood": "anxious",\n  "bodyState": ["疲惫", "睡不好"],\n  "note": "论文启动不了",\n  "shareWithChat": true\n}')
    add_bullet(doc, "能量固定为 1～5。")
    add_bullet(doc, "一天只有一条；再次保存同一天时覆盖旧记录。")
    add_bullet(doc, "浏览器本地最多保留最近 30 条记录。")
    add_bullet(doc, "不同情绪在历史列表中使用不同颜色显示。")
    add_bullet(doc, "当前没有数据库，因此不支持跨浏览器或跨设备同步。")
    add_bullet(doc, "关闭“分享给聊天助手”后，该条记录只留在本地。")

    add_heading(doc, "2.2 历史摘要的数据范围", 2)
    doc.add_paragraph("后端只分析最近 7 个自然日（包含今天）中允许分享的记录。未来日期、超过 7 天的日期和未分享记录都不进入 historyContext。")
    add_code(doc, "最近 7 天 = 今天、1 天前、2 天前、3 天前、4 天前、5 天前、6 天前")

    add_heading(doc, "2.3 加权能量均值", 2)
    doc.add_paragraph("能量均值不是简单平均。越接近今天的记录权重越高，避免很久之前的数据掩盖当前状态。")
    add_code(doc, "权重 = 1 / (1 + 距今天的天数 × 0.2)\n\n加权能量均值 = Σ(当天能量 × 当天权重) / Σ(当天权重)")
    add_bullet(doc, "今天：1.000")
    add_bullet(doc, "1 天前：0.833；2 天前：0.714；3 天前：0.625。")
    add_bullet(doc, "4 天前：0.556；5 天前：0.500；6 天前：0.455。")
    add_bullet(doc, "计算结果保留一位小数。")

    add_heading(doc, "2.4 低能量与趋势规则", 2)
    add_bullet(doc, "energy ≤ 2 记作一个低能量日。")
    add_bullet(doc, "趋势比较最近段（今天至 2 天前）与较早段（3 至 6 天前）的加权均值。")
    add_bullet(doc, "两段都至少有两条记录才计算趋势；否则为 insufficient_data。")
    add_bullet(doc, "最近段均值 - 较早段均值 ≥ 0.6：up。")
    add_bullet(doc, "最近段均值 - 较早段均值 ≤ -0.6：down。")
    add_bullet(doc, "其他情况：stable。")

    add_heading(doc, "2.5 数据覆盖、情绪和身体状态", 2)
    add_bullet(doc, "最近 7 天有效记录 0～2 条：coverage = insufficient。")
    add_bullet(doc, "3～5 条：coverage = partial。")
    add_bullet(doc, "6～7 条：coverage = good。")
    add_bullet(doc, "情绪按出现次数统计；并列最高的情绪会同时保留。")
    add_bullet(doc, "身体状态按出现次数排序，只取前三项 bodyStateTop。")

    add_heading(doc, "3. 前端、后端与平台的数据流", 1)
    add_code(doc, "周期页保存原始记录（本地）\n→ 前端筛选 shareWithChat = true 的最多 30 条\n→ 后端 buildHistoryContext(records)\n→ 计算最近 7 天的 historyContext\n→ 调用 OpenTrek 时传入工作流\n→ 输入标准化节点解析 historyContext\n→ 个性化回复节点按需参考")
    add_callout(doc, "隐私原则", "后端不持久保存这批记录；平台得到的是近期摘要，而不是完整的 30 天原始记录。用户当前原话始终优先于历史摘要。")

    add_heading(doc, "3.1 historyContext 示例", 2)
    add_code(doc, '{\n  "windowDays": 7,\n  "recordCount": 5,\n  "coverage": "partial",\n  "latestCheckinDate": "2026-07-31",\n  "energy": {\n    "recentAverage": 2.3,\n    "lowEnergyDays": 3,\n    "trend": "down"\n  },\n  "mood": {\n    "frequent": ["anxious", "irritable"]\n  },\n  "bodyStateTop": ["疲惫", "睡不好"]\n}')
    doc.add_paragraph("模型只能在确实有助于个性化建议时使用该摘要，不能报出分数、作医学诊断或把历史趋势当作因果证据。")

    add_heading(doc, "4. 历史上下文调试", 1)
    add_number(doc, "在前端“周期状态”页保存今天的记录，并开启“分享给聊天助手”。")
    add_number(doc, "返回聊天页，输入“我今天完全启动不了”。")
    add_number(doc, "进入平台 Trace，点击“输入标准化”节点。")
    add_number(doc, "确认输出中存在 hasHistoryContext = true 和 historyContext JSON。")
    add_number(doc, "点击“任务降级回复”或“情绪支持回复”，确认输入参数中含 historyContext。")
    add_number(doc, "检查回复是否更倾向低门槛、低能量友好的建议，但没有直接报出均值、记录数或医学结论。")
    add_code(doc, "hasHistoryContext: true\nhistoryContext: {\"windowDays\":7,...}")

    add_heading(doc, "5. 当前边界与后续方向", 1)
    add_bullet(doc, "已完成：呼吸训练、多轮确认、每日记录、本地历史、趋势摘要和平台 historyContext 接入。")
    add_bullet(doc, "尚未做：数据库、登录、跨设备同步、长期数据导出与删除。")
    add_bullet(doc, "尚未做：Agent 主动邀请“是否记录今天状态”；该能力可复用 pending_action 机制，后续再接入。")
    add_bullet(doc, "后续可将周期状态、今日记录和 historyContext 合并为统一的 personalContext，减少未来节点接入成本。")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
